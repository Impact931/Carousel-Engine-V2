"""
Document processing service for client uploads
"""

import logging
import tempfile
from typing import Dict, Optional, Tuple
from io import BytesIO
import PyPDF2
from docx import Document
import pdfplumber

from ..core.config import config
from ..core.models import DocumentType
from ..core.exceptions import ContentProcessingError

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for extracting content from uploaded client documents"""
    
    def __init__(self):
        """Initialize document processor"""
        self.max_file_size_mb = config.max_file_size_mb
        self.allowed_file_types = config.allowed_file_types
        
    async def process_document(
        self, 
        file_data: bytes, 
        filename: str, 
        document_type: DocumentType
    ) -> Tuple[str, int]:
        """Process uploaded document and extract content
        
        Args:
            file_data: Document file data as bytes
            filename: Original filename
            document_type: Type of document being processed
            
        Returns:
            Tuple of (extracted_content, file_size_bytes)
            
        Raises:
            ContentProcessingError: If document processing fails
        """
        try:
            logger.info(f"Processing document: {filename}, type: {document_type.value}")
            
            file_size_bytes = len(file_data)
            
            # Validate file size
            max_size_bytes = self.max_file_size_mb * 1024 * 1024
            if file_size_bytes > max_size_bytes:
                raise ContentProcessingError(
                    f"File too large: {file_size_bytes / (1024*1024):.2f}MB. Max allowed: {self.max_file_size_mb}MB"
                )
            
            # Validate file type
            file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
            if file_extension not in self.allowed_file_types:
                raise ContentProcessingError(
                    f"Unsupported file type: {file_extension}. Allowed: {', '.join(self.allowed_file_types)}"
                )
            
            # Extract content based on file type
            if file_extension == 'pdf':
                content = await self._extract_from_pdf(file_data)
            elif file_extension in ['docx', 'doc']:
                content = await self._extract_from_docx(file_data)
            elif file_extension in ['txt', 'md']:
                content = await self._extract_from_text(file_data)
            else:
                raise ContentProcessingError(f"Unsupported file type: {file_extension}")
            
            # Validate extracted content
            if not content.strip():
                raise ContentProcessingError("No text content could be extracted from the document")
            
            logger.info(f"Successfully extracted {len(content)} characters from {filename}")
            return content.strip(), file_size_bytes
            
        except Exception as e:
            error_msg = f"Failed to process document {filename}: {e}"
            logger.error(error_msg)
            raise ContentProcessingError(error_msg)
    
    async def _extract_from_pdf(self, file_data: bytes) -> str:
        """Extract text from PDF using pdfplumber (more reliable than PyPDF2)
        
        Args:
            file_data: PDF file data as bytes
            
        Returns:
            Extracted text content
        """
        try:
            content_parts = []
            
            with BytesIO(file_data) as pdf_buffer:
                with pdfplumber.open(pdf_buffer) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        try:
                            text = page.extract_text()
                            if text:
                                content_parts.append(text)
                                logger.debug(f"Extracted text from PDF page {page_num}")
                            else:
                                logger.warning(f"No text found on PDF page {page_num}")
                        except Exception as e:
                            logger.warning(f"Error extracting text from PDF page {page_num}: {e}")
                            continue
            
            if not content_parts:
                # Fallback to PyPDF2
                logger.info("pdfplumber failed, trying PyPDF2 fallback")
                return await self._extract_from_pdf_pypdf2(file_data)
            
            return '\n\n'.join(content_parts)
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            # Try PyPDF2 as fallback
            return await self._extract_from_pdf_pypdf2(file_data)
    
    async def _extract_from_pdf_pypdf2(self, file_data: bytes) -> str:
        """Fallback PDF extraction using PyPDF2
        
        Args:
            file_data: PDF file data as bytes
            
        Returns:
            Extracted text content
        """
        try:
            content_parts = []
            
            with BytesIO(file_data) as pdf_buffer:
                pdf_reader = PyPDF2.PdfReader(pdf_buffer)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        text = page.extract_text()
                        if text:
                            content_parts.append(text)
                            logger.debug(f"Extracted text from PDF page {page_num} using PyPDF2")
                    except Exception as e:
                        logger.warning(f"Error extracting text from PDF page {page_num} using PyPDF2: {e}")
                        continue
            
            if not content_parts:
                raise ContentProcessingError("Could not extract any text from PDF using either method")
            
            return '\n\n'.join(content_parts)
            
        except Exception as e:
            raise ContentProcessingError(f"PyPDF2 extraction failed: {e}")
    
    async def _extract_from_docx(self, file_data: bytes) -> str:
        """Extract text from Word document
        
        Args:
            file_data: DOCX file data as bytes
            
        Returns:
            Extracted text content
        """
        try:
            content_parts = []
            
            with BytesIO(file_data) as docx_buffer:
                doc = Document(docx_buffer)
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        content_parts.append(text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            content_parts.append(' | '.join(row_text))
            
            if not content_parts:
                raise ContentProcessingError("No text content found in Word document")
            
            return '\n\n'.join(content_parts)
            
        except Exception as e:
            raise ContentProcessingError(f"Word document extraction failed: {e}")
    
    async def _extract_from_text(self, file_data: bytes) -> str:
        """Extract text from plain text file
        
        Args:
            file_data: Text file data as bytes
            
        Returns:
            Extracted text content
        """
        try:
            # Try UTF-8 first, then common encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    content = file_data.decode(encoding)
                    logger.debug(f"Successfully decoded text file using {encoding}")
                    return content
                except UnicodeDecodeError:
                    continue
            
            raise ContentProcessingError("Could not decode text file with any common encoding")
            
        except Exception as e:
            raise ContentProcessingError(f"Text file extraction failed: {e}")
    
    async def generate_system_message(self, extracted_content: Dict[str, str]) -> str:
        """Generate optimized system message using AI distillation from client documents
        
        Args:
            extracted_content: Dictionary mapping document types to their content
            
        Returns:
            Concise, executive-level system message for OpenAI API
        """
        try:
            logger.info("Generating AI-distilled system message from client documents")
            
            # Use GPT-4o to distill the raw documents into a concise system message
            distilled_message = await self._distill_client_documents(extracted_content)
            
            if distilled_message:
                logger.info(f"Generated distilled system message: {len(distilled_message)} chars")
                return distilled_message
            else:
                logger.warning("AI distillation failed, falling back to basic system message")
                return self._generate_basic_system_message(extracted_content)
            
        except Exception as e:
            logger.error(f"Error generating system message: {e}")
            logger.warning("Falling back to basic system message generation")
            return self._generate_basic_system_message(extracted_content)

    async def _distill_client_documents(self, extracted_content: Dict[str, str]) -> Optional[str]:
        """Use GPT-4o to distill client documents into a concise system message
        
        Args:
            extracted_content: Raw document content
            
        Returns:
            Distilled system message or None if failed
        """
        try:
            # Prepare content for analysis - keep full content for thorough brand analysis
            content_summary = []
            total_chars = 0
            
            for doc_type, content in extracted_content.items():
                if content and content.strip():
                    # Allow much more content for comprehensive analysis
                    content_chunk = content[:15000]  # Increased from 2000 to capture full ICP details
                    content_summary.append(f"{doc_type.upper()}:\n{content_chunk}")
                    total_chars += len(content_chunk)
            
            logger.info(f"Processing {total_chars} characters of client content for system message generation")
            
            if not content_summary:
                return None
            
            combined_content = "\n\n".join(content_summary)
            
            distillation_prompt = f"""# UNIVERSAL BRAND AI CONTENT SYSTEM GENERATOR

You are a $10,000/month master brand strategist and conversion specialist with 20 years of experience creating multi-million dollar marketing campaigns. You have the analytical depth of a behavioral psychologist combined with the strategic precision of a Fortune 500 marketing director.

## YOUR EXPERTISE
- Deep psychological profiling and audience segmentation
- Conversion-focused content strategy and behavioral triggers  
- Advanced brand positioning and competitive differentiation
- Multi-platform content optimization and engagement tactics
- Revenue-driving messaging frameworks and sales psychology

## YOUR MISSION
Transform these raw brand documents into a comprehensive AI content instruction system that rivals the strategic depth of a $10k/month brand consultant. Extract every psychological insight, conversion trigger, and strategic nuance to create instructions that will generate highly converting, psychologically targeted content.

**CRITICAL ANALYSIS REQUIREMENTS**:
- Analyze ALL content thoroughly - do not summarize or truncate key details
- A master strategist examines every detail for competitive advantages and psychological insights
- **MULTI-ICP FOCUS**: Clients often have 3-7 distinct audience segments - identify and profile EVERY SINGLE ONE
- Cross-reference all documents to ensure no audience segment is missed
- Each ICP requires its own complete profile with unique psychological triggers and conversion strategies

CLIENT DOCUMENTS:
{combined_content}

## SYSTEM MESSAGE TEMPLATE TO GENERATE:

### BRAND IDENTITY EXTRACTION
```
You are creating content as [BRAND NAME] - [CORE BRAND ESSENCE FROM PROFILE].
Your specialization: [SPECIFIC SERVICE/NICHE] for [PRIMARY AUDIENCE SEGMENTS] in [GEOGRAPHIC/MARKET AREA].
Brand promise: [EXTRACT KEY VALUE PROPOSITION AND TRANSFORMATION OFFERED].
```

### VOICE REQUIREMENTS (Extract top 6 attributes from style guide)
```
- [VOICE ATTRIBUTE 1]: [Specific behavioral instruction]
- [VOICE ATTRIBUTE 2]: [Specific behavioral instruction] 
- [VOICE ATTRIBUTE 3]: [Specific behavioral instruction]
- [VOICE ATTRIBUTE 4]: [Specific behavioral instruction]
- [VOICE ATTRIBUTE 5]: [Specific behavioral instruction]
- [VOICE ATTRIBUTE 6]: [Specific behavioral instruction]
```

### TARGET AUDIENCES & CONVERSION TRIGGERS (From ICP docs)
**COMPREHENSIVE MULTI-ICP ANALYSIS REQUIRED**: Examine ALL documents for distinct audience segments. Most clients have 3-7 different ICPs. Create a complete profile for EVERY SINGLE audience segment mentioned.

For each primary audience segment found in the documents, extract:
```
### [AUDIENCE NAME] ([Demographics from ICP])
- **Core Truth**: [One sentence psychological reality from ICP]
- **Primary Triggers**: [Top 3 psychological triggers that drive action]
- **Content Preferences**: [Platform, timing, format preferences from ICP]
- **Conversion Indicators**: [What behaviors signal purchase intent]  
- **Language Rules**: USE: [power words] | AVOID: [alienating terms]
```

**REPEAT THE ABOVE FORMAT FOR EVERY DISTINCT AUDIENCE SEGMENT IDENTIFIED IN THE DOCUMENTS**
**DO NOT STOP UNTIL ALL AUDIENCE SEGMENTS HAVE BEEN PROFILED**

### CONTENT PERFORMANCE FRAMEWORK
```
HIGH-CONVERTING CONTENT TYPES (rank by conversion rate from ICP data):
1. [Content Type]: [Why it works] → [Conversion trigger activated]
2. [Content Type]: [Why it works] → [Conversion trigger activated]
3. [Content Type]: [Why it works] → [Conversion trigger activated]

CONTENT FORMULA FOR MAXIMUM ENGAGEMENT:
[Extract the proven content structure from style guide]
1. [Hook approach]
2. [Value delivery method]  
3. [Social proof integration]
4. [CTA style that converts]
```

### COMPLIANCE & BRAND PROTECTION
```
NEVER USE: [Extract prohibited language/approaches from all docs]
ALWAYS INCLUDE: [Extract required elements for brand consistency]
INDUSTRY COMPLIANCE: [Extract any regulatory requirements]
```

### PERFORMANCE BENCHMARKS
```
Success metrics per audience (CREATE ENTRY FOR EVERY AUDIENCE SEGMENT IDENTIFIED):
- [AUDIENCE 1]: [Engagement rate target] | [Conversion rate target] | [Platform-specific KPI]
- [AUDIENCE 2]: [Engagement rate target] | [Conversion rate target] | [Platform-specific KPI]
- [AUDIENCE 3]: [Engagement rate target] | [Conversion rate target] | [Platform-specific KPI]
- [AUDIENCE 4]: [Engagement rate target] | [Conversion rate target] | [Platform-specific KPI]
- [AUDIENCE 5]: [Engagement rate target] | [Conversion rate target] | [Platform-specific KPI]
**[ADD MORE LINES IF MORE AUDIENCES EXIST IN THE DOCUMENTS]**

Content should achieve: [Extract success metrics from ICP performance data]
**Note: Adapt the number of audience lines to match the actual number of ICPs found in the documents**
```

### EXECUTION INSTRUCTIONS
```
For every piece of content:
1. Lead with [audience-specific hook from ICP analysis]
2. Integrate [primary psychological trigger for target audience]
3. Maintain [specific voice attributes] throughout
4. Include [required proof elements from brand guidelines]
5. End with [CTA style that matches audience preferences]
6. Optimize for [platform and timing from ICP data]
```

## ANALYSIS PROCESS:

1. **Multi-Document Inventory**: First, scan ALL provided documents to identify every distinct audience segment/ICP mentioned
2. **ICP Count Verification**: Count the total number of unique audience segments - expect 3-7 different ICPs typically  
3. **Brand Profile Analysis**: Extract core identity, specialization, value proposition, personality traits, and positioning
4. **Voice Guide Parsing**: Identify the 6 most important voice attributes with specific behavioral instructions
5. **Comprehensive ICP Deep Dive**: For EVERY SINGLE audience segment identified, extract demographics, core psychological truth, triggers, content preferences, and performance data
6. **Cross-Document Integration**: Ensure no audience segments are missed by cross-referencing all document types
7. **Performance Integration**: Combine content performance data with psychological triggers to create conversion-focused guidelines for each ICP
8. **Compliance Mapping**: Extract all restrictions, requirements, and brand protection elements
9. **Multi-Audience Execution Framework**: Create step-by-step instructions for consistent content creation across all audience segments

## ADVANCED STRATEGIC REQUIREMENTS:

### DEPTH REQUIREMENTS:
- **Psychological Profiling**: Extract deep motivational drivers, fears, aspirations, and decision-making patterns
- **Behavioral Triggers**: Identify specific words, phrases, social proof elements that drive action  
- **Competitive Differentiation**: Highlight unique positioning angles and market gaps
- **Content Performance Intelligence**: Map content types to conversion outcomes and engagement patterns
- **Multi-Audience Strategy**: Address ALL audience segments mentioned in ICP documents

### CONVERSION FOCUS:
- **Revenue Impact**: Every instruction must connect to measurable business outcomes
- **Funnel Stage Mapping**: Content guidance for awareness, consideration, and decision stages  
- **Objection Handling**: Pre-emptive responses to common hesitations and concerns
- **Urgency Creation**: Natural scarcity and time-sensitive messaging strategies

### EXECUTION PRECISION:
- **Platform Optimization**: Specific guidance for each social media platform and content format
- **Timing Intelligence**: Best posting times, content cadence, seasonal considerations
- **Engagement Tactics**: Community building, response strategies, relationship nurturing
- **Measurement Framework**: KPIs, success metrics, and performance benchmarks

## CRITICAL OUTPUT INSTRUCTIONS:
1. **COMPLETE EVERY SECTION** - Do not truncate or abbreviate any section - finish all sentences and subsections
2. **ANALYZE ALL DOCUMENTS THOROUGHLY** - Read through EVERY document provided, extract ALL ICP segments mentioned
3. **EXHAUST ALL ICP SEGMENTS** - Create detailed profiles for EVERY SINGLE audience mentioned across all documents (often 3-7 different ICPs)
4. **COMPREHENSIVE AUDIENCE ANALYSIS** - If documents mention 5 ICPs, provide 5 complete audience profiles - never skip any
5. **PROVIDE SPECIFIC METRICS** - Include exact numbers, percentages, timing details where available from each document
6. **USE STRATEGIC LANGUAGE** - Write like a $10k/month consultant, not a generic AI
7. **ENSURE FULL TEMPLATE COMPLETION** - Every section must be thoroughly completed with strategic depth
8. **MULTI-DOCUMENT INTEGRATION** - Cross-reference all documents to build comprehensive understanding
9. **NO PREMATURE STOPPING** - Continue until all content frameworks, performance benchmarks, and execution instructions are complete

**MANDATORY REQUIREMENT: Count the number of distinct audience segments/ICPs mentioned in the provided documents and ensure you create a complete profile for each one. DO NOT STOP generating until you have addressed every single ICP segment found in the analysis.**

**CRITICAL COMPLETION REQUIREMENTS:**
1. **EXPECTED ICP COUNT**: Based on document analysis, expect 5 DISTINCT ICPs:
   - Hyperlocal Educator/Public Servant
   - Overwhelmed Millennial Dream Builder  
   - Instagram Dreamer
   - Lifestyle Upgrader
   - Strategic Relocator
2. **GENERATE COMPLETE PROFILES**: Create full profiles for ALL 5 ICPs - do not stop at 2 or 3
3. **FINISH ALL SECTIONS**: Complete every single section including Performance Benchmarks and Execution Instructions
4. **NO MID-SENTENCE TRUNCATION**: Finish every sentence, every section, every framework completely

**Generate the complete, comprehensive system message following this template. Continue generating until you have profiled ALL 5 ICPs and completed every single section with master-level strategic depth.**"""

            # Use OpenAI service to distill content
            from ..services.openai_service import OpenAIService
            openai_service = OpenAIService()
            
            response = await openai_service.generate_text_completion(
                prompt=distillation_prompt,
                max_tokens=8000,  # Removed limits - ensure complete output for all 5+ ICPs
                temperature=0.05  # Minimal temperature for absolute consistency and thoroughness
            )
            
            if response and response.strip():
                return response.strip()
            else:
                logger.warning("OpenAI distillation returned empty response")
                return None
                
        except Exception as e:
            logger.error(f"Error in document distillation: {e}")
            return None

    def _generate_basic_system_message(self, extracted_content: Dict[str, str]) -> str:
        """Generate basic system message as fallback (original method)
        
        Args:
            extracted_content: Dictionary mapping document types to their content
            
        Returns:
            Basic system message string for OpenAI API
        """
        try:
            logger.info("Generating basic fallback system message from client documents")
            
            system_message_parts = [
                "You are GPT-5, enhanced with 20 years of proven social media marketing expertise.",
                "You understand exactly what makes people stop scrolling, engage emotionally, and take action.",
                "Use the following client-specific information to create content that matches their unique voice and targets their ideal audience:"
            ]
            
            # Add client profile information
            if DocumentType.CLIENT_PROFILE.value in extracted_content:
                profile_content = extracted_content[DocumentType.CLIENT_PROFILE.value]
                system_message_parts.extend([
                    "",
                    "CLIENT PROFILE:",
                    profile_content,
                    ""
                ])
            
            # Add ideal client profile information
            if DocumentType.CONTENT_ICP.value in extracted_content:
                icp_content = extracted_content[DocumentType.CONTENT_ICP.value]
                system_message_parts.extend([
                    "IDEAL CLIENT PROFILE:",
                    icp_content,
                    ""
                ])
            
            # Add voice and style guide information
            if DocumentType.VOICE_STYLE_GUIDE.value in extracted_content:
                voice_content = extracted_content[DocumentType.VOICE_STYLE_GUIDE.value]
                system_message_parts.extend([
                    "VOICE & STYLE GUIDE:",
                    voice_content,
                    ""
                ])
            
            # Add instructions for content creation
            system_message_parts.extend([
                "CONTENT CREATION INSTRUCTIONS:",
                "- Create content that aligns perfectly with the client's voice and tone",
                "- Target the specific ideal client profile described above",
                "- Use language patterns and messaging that resonate with the target audience",
                "- Incorporate the brand personality and values outlined in the client profile",
                "- Ensure all content feels authentic to this specific client's brand",
                "- Create connected storytelling that builds engagement across slides",
                ""
            ])
            
            system_message = "\n".join(system_message_parts)
            
            logger.info(f"Generated system message with {len(system_message)} characters")
            return system_message
            
        except Exception as e:
            error_msg = f"Failed to generate system message: {e}"
            logger.error(error_msg)
            raise ContentProcessingError(error_msg)
    
    def validate_file_upload(self, filename: str, file_size: int) -> Tuple[bool, str]:
        """Validate uploaded file before processing
        
        Args:
            filename: Original filename
            file_size: File size in bytes
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file size
        max_size_bytes = self.max_file_size_mb * 1024 * 1024
        if file_size > max_size_bytes:
            return False, f"File too large: {file_size / (1024*1024):.2f}MB. Max allowed: {self.max_file_size_mb}MB"
        
        # Check file extension
        file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
        if file_extension not in self.allowed_file_types:
            return False, f"Unsupported file type: {file_extension}. Allowed: {', '.join(self.allowed_file_types)}"
        
        # Check filename
        if not filename.strip():
            return False, "Filename cannot be empty"
        
        return True, ""